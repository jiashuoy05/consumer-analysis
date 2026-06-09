import os
import smtplib
import tempfile
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from fpdf import FPDF

from dotenv import load_dotenv

load_dotenv(Path(__file__).parent.parent / ".env")

SMTP_HOST = os.getenv("SMTP_HOST", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASS = os.getenv("SMTP_PASS", "")

CHART_COLORS = ["#FF6384", "#36A2EB", "#FFCE56", "#4BC0C0", "#9966FF", "#FF9F40"]
CATEGORY_LABELS = {"食": "飲食", "衣": "衣著", "住": "居住", "行": "交通", "育": "教育", "樂": "娛樂"}

_OUTPUT_DIR = Path(tempfile.gettempdir()) / "consumer-analysis-reports"
_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

_FONT_CANDIDATES = [
    "C:/Windows/Fonts/kaiu.ttf",
    "C:/Windows/Fonts/simsunb.ttf",
    "C:/Windows/Fonts/SimsunExtG.ttf",
]


def _find_chinese_font() -> str | None:
    for p in _FONT_CANDIDATES:
        if Path(p).exists():
            return p
    return None


def _get_chart_font() -> FontProperties | None:
    fp = _find_chinese_font()
    if fp:
        return FontProperties(fname=fp, size=12)
    return None


def _sum_by_category(classified: list[dict]) -> dict:
    cat_totals: dict[str, float] = {}
    for exp in classified:
        cat = exp.get("category", "其他")
        cat_totals[cat] = cat_totals.get(cat, 0) + float(exp.get("amount", 0))
    return cat_totals


def generate_pie_chart(classified: list[dict]) -> str:
    cat_totals = _sum_by_category(classified)
    if not cat_totals:
        return ""

    sorted_cats = sorted(cat_totals.items(), key=lambda x: x[1], reverse=True)
    labels = [CATEGORY_LABELS.get(c, c) for c, _ in sorted_cats]
    amounts = [v for _, v in sorted_cats]
    colors = CHART_COLORS[:len(labels)]

    fig, ax = plt.subplots(figsize=(6, 5))
    font = _get_chart_font()

    wedges, texts, autotexts = ax.pie(
        amounts,
        labels=labels,
        autopct="%1.1f%%",
        colors=colors,
        startangle=90,
        textprops={"fontproperties": font} if font else {},
    )
    for t in autotexts:
        t.set_fontsize(11)
        if font:
            t.set_fontproperties(font)

    ax.set_title("消費類別佔比分析", fontsize=14, fontweight="bold",
                 fontproperties=font if font else None)

    chart_path = str(_OUTPUT_DIR / "pie_chart.png")
    plt.savefig(chart_path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return chart_path


def create_report_docx(
    report: dict,
    classified: list[dict],
    answers: dict[str, str],
    chart_path: str,
    period: str = "",
) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(11)

    h1_style = doc.styles["Heading 1"]
    h1_style.font.name = "Microsoft JhengHei"
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x3D)

    h2_style = doc.styles["Heading 2"]
    h2_style.font.name = "Microsoft JhengHei"
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x4B)

    cat_totals = _sum_by_category(classified)

    doc.add_heading(f"AI 消費幸福感分析報告（{period}）" if period else "AI 消費幸福感分析報告", level=1)
    doc.add_paragraph("")

    doc.add_heading("一、消費總覽", level=2)
    total = sum(cat_totals.values())
    p = doc.add_paragraph()
    p.add_run(f"本期總消費金額：NT$ {total:,.0f}").bold = True
    p.add_run(f"\n涵蓋 {len(classified)} 筆明細，分為 {len(cat_totals)} 大類別。")

    for cat, amt in sorted(cat_totals.items(), key=lambda x: x[1], reverse=True):
        label = CATEGORY_LABELS.get(cat, cat)
        doc.add_paragraph(f"  {label}：NT$ {amt:,.0f}（佔比 {amt/total*100:.1f}%）")

    doc.add_paragraph("")

    if chart_path and os.path.exists(chart_path):
        doc.add_heading("二、消費類別圖表", level=2)
        doc.add_picture(chart_path, width=Inches(5))

    doc.add_paragraph("")

    doc.add_heading("三、幸福消費 Top 3", level=2)
    for i, item in enumerate(report.get("happy_top3", []), 1):
        doc.add_paragraph(f"#{i}  {item}")

    doc.add_paragraph("")

    doc.add_heading("四、壓力消費 Top 3", level=2)
    for i, item in enumerate(report.get("stress_top3", []), 1):
        doc.add_paragraph(f"#{i}  {item}")

    doc.add_paragraph("")

    doc.add_heading("五、改善建議", level=2)
    for i, item in enumerate(report.get("suggestions", []), 1):
        doc.add_paragraph(f"{i}. {item}")

    doc.add_paragraph("")

    doc.add_heading("六、整體總結", level=2)
    doc.add_paragraph(report.get("summary", ""))

    doc.add_paragraph("")

    doc.add_heading("七、問卷回饋", level=2)
    for q, a in answers.items():
        p = doc.add_paragraph()
        run_q = p.add_run("問題：")
        run_q.bold = True
        p.add_run(q)
        p2 = doc.add_paragraph()
        run_a = p2.add_run("回答：")
        run_a.bold = True
        p2.add_run(a)
        doc.add_paragraph("")

    output_path = str(_OUTPUT_DIR / "happiness_report.docx")
    doc.save(output_path)
    return output_path


def create_report_pdf(
    report: dict,
    classified: list[dict],
    answers: dict[str, str],
    chart_path: str,
    period: str = "",
) -> str:
    cjk_font = _find_chinese_font()
    if not cjk_font:
        raise RuntimeError("找不到中文字型，無法生成 PDF")

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.add_page()

    pdf.add_font("TNR", "", "C:/Windows/Fonts/times.ttf")
    pdf.add_font("TNR", "B", "C:/Windows/Fonts/timesbd.ttf")
    pdf.add_font("TNR", "I", "C:/Windows/Fonts/timesi.ttf")
    pdf.add_font("TNR", "BI", "C:/Windows/Fonts/timesbi.ttf")
    pdf.add_font("KAI", "", cjk_font)
    pdf.set_fallback_fonts(["KAI"])

    def heading(text: str, size: int = 16):
        pdf.set_font("TNR", "B", size)
        pdf.cell(0, 12, text, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    def body(text: str, size: int = 11):
        pdf.set_font("TNR", "", size)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(pdf.epw, 7, text)
        pdf.ln(1)

    def body_bold_prefix(prefix: str, text: str, size: int = 11):
        pdf.set_font("TNR", "B", size)
        pdf.set_x(pdf.l_margin)
        full = prefix + text
        pdf.multi_cell(pdf.epw, 7, full)

    pdf.set_font("TNR", "B", 24)
    title = f"AI 消費幸福感分析報告（{period}）" if period else "AI 消費幸福感分析報告"
    pdf.cell(0, 15, title, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    cat_totals = _sum_by_category(classified)
    total = sum(cat_totals.values())

    heading("一、消費總覽", 14)
    body(f"本期總消費金額：NT$ {total:,.0f}")
    body(f"涵蓋 {len(classified)} 筆明細，分為 {len(cat_totals)} 大類別。")
    pdf.ln(2)
    for cat, amt in sorted(cat_totals.items(), key=lambda x: x[1], reverse=True):
        label = CATEGORY_LABELS.get(cat, cat)
        body(f"  {label}：NT$ {amt:,.0f}（佔比 {amt/total*100:.1f}%）")

    pdf.ln(4)

    if chart_path and os.path.exists(chart_path):
        heading("二、消費類別圖表", 14)
        pdf.image(chart_path, x=30, w=150)
        pdf.ln(4)

    heading("三、幸福消費 Top 3", 14)
    for i, item in enumerate(report.get("happy_top3", []), 1):
        body(f"#{i}  {item}")
    pdf.ln(3)

    heading("四、壓力消費 Top 3", 14)
    for i, item in enumerate(report.get("stress_top3", []), 1):
        body(f"#{i}  {item}")
    pdf.ln(3)

    heading("五、改善建議", 14)
    for i, item in enumerate(report.get("suggestions", []), 1):
        body(f"{i}. {item}")
    pdf.ln(3)

    heading("六、整體總結", 14)
    body(report.get("summary", ""))
    pdf.ln(4)

    heading("七、問卷回饋", 14)
    for q, a in answers.items():
        body_bold_prefix("問題：", q)
        body_bold_prefix("回答：", a)
        pdf.ln(2)

    output_path = str(_OUTPUT_DIR / "happiness_report.pdf")
    pdf.output(output_path)
    return output_path


def send_email(to_email: str, pdf_path: str, smtp_config: dict | None = None, period: str = "") -> None:
    host = smtp_config.get("host", SMTP_HOST) if smtp_config else SMTP_HOST
    port = smtp_config.get("port", SMTP_PORT) if smtp_config else SMTP_PORT
    user = smtp_config.get("user", SMTP_USER) if smtp_config else SMTP_USER
    password = smtp_config.get("pass", SMTP_PASS) if smtp_config else SMTP_PASS

    if not user or not password:
        raise RuntimeError("SMTP 未設定，請檢查 SMTP_USER / SMTP_PASS 環境變數")

    label = f"（{period}）" if period else ""

    msg = MIMEMultipart()
    msg["From"] = user
    msg["To"] = to_email
    msg["Subject"] = f"AI 消費幸福感分析報告{label}"

    body = f"""\
<html>
<head><meta charset="utf-8"></head>
<body style="font-family: 'Segoe UI', Arial, sans-serif; max-width: 560px; margin: 0 auto; padding: 20px; color: #333;">
  <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 8px; padding: 24px; text-align: center; margin-bottom: 20px;">
    <h1 style="color: #fff; margin: 0; font-size: 22px;">📊 AI 消費幸福感分析報告</h1>
  </div>

  <p style="font-size: 15px;">您好，</p>

  <p style="font-size: 15px;">附件為您的 AI 消費幸福感分析報告（PDF）{label}，包含以下內容：</p>

  <ul style="font-size: 14px; line-height: 1.8; padding-left: 20px;">
    <li>📈 <strong>消費分類總覽</strong> — 食/衣/住/行/育/樂 分佈圖表</li>
    <li>😊 <strong>幸福消費 Top 3</strong> — 讓您感到快樂的支出</li>
    <li>😰 <strong>壓力消費 Top 3</strong> — 高金額低滿意度的支出</li>
    <li>💡 <strong>個人化改善建議</strong> — 下個月的財務小目標</li>
    <li>📝 <strong>問卷回饋內容</strong> — 您的消費心情記錄</li>
  </ul>

  <p style="font-size: 14px; color: #888;">感謝您的使用，我們下個月見！</p>

  <div style="border-top: 1px solid #eee; margin-top: 24px; padding-top: 12px; font-size: 12px; color: #aaa; text-align: center;">
    AI 消費幸福感分析師
  </div>
</body>
</html>
"""
    msg.attach(MIMEText(body, "html", "utf-8"))

    filename = f"happiness_report_{period}.pdf" if period else "happiness_report.pdf"
    with open(pdf_path, "rb") as f:
        part = MIMEBase("application", "pdf")
        part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header(
            "Content-Disposition",
            "attachment",
            filename=filename,
        )
        msg.attach(part)

    with smtplib.SMTP(host, port) as server:
        server.starttls()
        server.login(user, password)
        server.send_message(msg)


def cleanup(temp_path: str | None = None) -> None:
    if temp_path and os.path.exists(temp_path):
        os.remove(temp_path)

    chart = _OUTPUT_DIR / "pie_chart.png"
    if chart.exists():
        os.remove(chart)
